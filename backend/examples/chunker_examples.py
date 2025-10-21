"""
Examples demonstrating the PrivAI FileChunker functionality
"""
import json
import tempfile
from pathlib import Path
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Inches

from app.core.chunker import FileChunker, parse_file_to_chunks, parse_multiple_files


def create_sample_pdf():
    """Create a sample PDF file for testing (simulated with text file)"""
    content = """
    PRIVAI SAMPLE DOCUMENT
    =====================
    
    This is a comprehensive sample document for testing the PrivAI chunker.
    It contains multiple sections with various types of content.
    
    SECTION 1: INTRODUCTION
    ----------------------
    PrivAI is a privacy-first AI application designed for college data processing.
    It provides secure, local processing of documents without sending data to external services.
    
    Key Features:
    - Local file processing
    - Vector-based document search
    - AI-powered question answering
    - Privacy-first architecture
    
    SECTION 2: TECHNICAL SPECIFICATIONS
    ----------------------------------
    The application uses advanced natural language processing techniques including:
    - Text chunking with intelligent overlap
    - Vector embeddings using sentence transformers
    - FAISS vector database for similarity search
    - Local LLM integration for responses
    
    Performance Metrics:
    - Chunk size: 500 tokens (configurable)
    - Overlap: 50 tokens (configurable)
    - Processing speed: ~1000 tokens/second
    - Memory usage: Optimized for local deployment
    
    SECTION 3: USE CASES
    -------------------
    PrivAI is designed for educational institutions and includes:
    
    1. Student Data Processing
       - Academic records analysis
       - Course material indexing
       - Research paper processing
    
    2. Administrative Tasks
       - Document search and retrieval
       - Policy document analysis
       - Compliance checking
    
    3. Research Applications
       - Literature review assistance
       - Data extraction from papers
       - Knowledge base creation
    
    SECTION 4: IMPLEMENTATION DETAILS
    --------------------------------
    The chunker module provides:
    - Multi-format file support (PDF, DOCX, CSV, Excel)
    - Intelligent text splitting
    - Metadata preservation
    - Token counting and optimization
    
    Each chunk includes comprehensive metadata:
    - Source file information
    - Page/row numbers
    - Timestamp
    - Token count
    - Chunk position
    
    CONCLUSION
    ----------
    PrivAI represents a significant advancement in privacy-preserving AI applications
    for educational institutions. By keeping all processing local and providing
    comprehensive metadata tracking, it ensures both functionality and privacy.
    
    For more information, visit the PrivAI documentation or contact the development team.
    """
    
    return content


def create_sample_docx():
    """Create a sample DOCX file for testing"""
    doc = Document()
    
    # Title
    title = doc.add_heading('PrivAI Technical Documentation', 0)
    
    # Introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'PrivAI is a comprehensive AI platform designed for educational institutions. '
        'It provides secure, local processing of academic documents and data.'
    )
    
    # Features section
    doc.add_heading('Key Features', level=1)
    features = [
        'Local file processing and analysis',
        'Vector-based document search',
        'AI-powered question answering',
        'Multi-format file support',
        'Privacy-first architecture',
        'Real-time document indexing'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Add a table
    doc.add_heading('Supported File Types', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'File Type'
    hdr_cells[1].text = 'Extension'
    hdr_cells[2].text = 'Parser Used'
    
    file_types = [
        ('PDF Documents', '.pdf', 'pdfplumber'),
        ('Word Documents', '.docx', 'python-docx'),
        ('Spreadsheets', '.csv, .xlsx', 'pandas'),
        ('Text Files', '.txt', 'built-in')
    ]
    
    for file_type, extension, parser in file_types:
        row_cells = table.add_row().cells
        row_cells[0].text = file_type
        row_cells[1].text = extension
        row_cells[2].text = parser
    
    # Technical details
    doc.add_heading('Technical Implementation', level=1)
    doc.add_paragraph(
        'The PrivAI chunker uses advanced natural language processing techniques '
        'to intelligently split documents into meaningful chunks while preserving '
        'context and metadata.'
    )
    
    doc.add_paragraph(
        'Key technical features include intelligent sentence boundary detection, '
        'token counting optimization, and comprehensive metadata tracking for '
        'each generated chunk.'
    )
    
    return doc


def create_sample_csv():
    """Create a sample CSV file for testing"""
    data = {
        'Student_ID': [1001, 1002, 1003, 1004, 1005],
        'Name': ['Alice Johnson', 'Bob Smith', 'Carol Davis', 'David Wilson', 'Eva Brown'],
        'Major': ['Computer Science', 'Mathematics', 'Physics', 'Biology', 'Chemistry'],
        'GPA': [3.8, 3.6, 3.9, 3.7, 3.5],
        'Year': ['Senior', 'Junior', 'Senior', 'Sophomore', 'Junior'],
        'Courses': [
            'CS401, CS402, CS403',
            'MATH301, MATH302, MATH303',
            'PHYS401, PHYS402, PHYS403',
            'BIO201, BIO202, BIO203',
            'CHEM301, CHEM302, CHEM303'
        ]
    }
    
    df = pd.DataFrame(data)
    return df


def demonstrate_chunker():
    """Demonstrate the chunker with different file types"""
    print("🚀 PrivAI FileChunker Demonstration")
    print("=" * 60)
    
    # Initialize chunker with custom settings
    chunker = FileChunker(chunk_size=500, chunk_overlap=50)
    print(f"✅ Chunker initialized: {chunker.chunk_size} tokens, {chunker.chunk_overlap} overlap")
    
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        
        # Test 1: Text file (simulating PDF)
        print(f"\n📄 Test 1: Text Document (PDF simulation)")
        print("-" * 40)
        
        text_file = temp_path / "sample.txt"
        with open(text_file, "w", encoding="utf-8") as f:
            f.write(create_sample_pdf())
        
        chunks = chunker.parse_file(text_file)
        print(f"✅ Created {len(chunks)} chunks from text document")
        
        # Show first chunk details
        if chunks:
            chunk = chunks[0]
            print(f"   First chunk preview: {chunk['text'][:100]}...")
            print(f"   Metadata: {chunk['metadata']['token_count']} tokens, "
                  f"chunk {chunk['metadata']['chunk_index'] + 1}")
        
        # Test 2: DOCX file
        print(f"\n📝 Test 2: Word Document (DOCX)")
        print("-" * 40)
        
        docx_file = temp_path / "sample.docx"
        doc = create_sample_docx()
        doc.save(docx_file)
        
        chunks = chunker.parse_file(docx_file)
        print(f"✅ Created {len(chunks)} chunks from DOCX document")
        
        # Show chunk with table data
        for chunk in chunks:
            if "Table:" in chunk['text']:
                print(f"   Table chunk found: {chunk['text'][:100]}...")
                break
        
        # Test 3: CSV file
        print(f"\n📊 Test 3: CSV Data")
        print("-" * 40)
        
        csv_file = temp_path / "sample.csv"
        df = create_sample_csv()
        df.to_csv(csv_file, index=False)
        
        chunks = chunker.parse_file(csv_file)
        print(f"✅ Created {len(chunks)} chunks from CSV data")
        
        # Show data chunk
        if chunks:
            chunk = chunks[0]
            print(f"   Data chunk preview: {chunk['text'][:150]}...")
            print(f"   Contains {chunk['metadata']['token_count']} tokens")
        
        # Test 4: Multiple files
        print(f"\n📁 Test 4: Multiple Files Processing")
        print("-" * 40)
        
        all_chunks = parse_multiple_files([text_file, docx_file, csv_file])
        print(f"✅ Processed {len(all_chunks)} total chunks from 3 files")
        
        # Group by source
        sources = {}
        for chunk in all_chunks:
            source = chunk['metadata']['source_file']
            if source not in sources:
                sources[source] = 0
            sources[source] += 1
        
        print(f"   Chunks per file:")
        for source, count in sources.items():
            print(f"     {Path(source).name}: {count} chunks")
        
        # Test 5: Statistics
        print(f"\n📊 Test 5: Chunking Statistics")
        print("-" * 40)
        
        stats = chunker.get_chunking_stats(all_chunks)
        print(f"   Total chunks: {stats['total_chunks']}")
        print(f"   Average tokens per chunk: {stats['avg_tokens_per_chunk']:.1f}")
        print(f"   Token range: {stats['min_tokens']} - {stats['max_tokens']}")
        print(f"   Total tokens: {stats['total_tokens']}")
        
        # Test 6: JSON serialization
        print(f"\n💾 Test 6: JSON Serialization")
        print("-" * 40)
        
        sample_chunk = all_chunks[0]
        json_str = json.dumps(sample_chunk, indent=2)
        print(f"✅ JSON serialization successful")
        print(f"   Sample chunk JSON length: {len(json_str)} characters")
        
        # Show metadata structure
        print(f"\n🏷️  Metadata Structure:")
        print(f"   {list(sample_chunk['metadata'].keys())}")
    
    print(f"\n🎉 All demonstrations completed successfully!")


def show_chunker_capabilities():
    """Show the capabilities of the chunker"""
    print(f"\n🔧 PrivAI FileChunker Capabilities")
    print("=" * 60)
    
    chunker = FileChunker()
    
    print(f"📁 Supported File Types:")
    for ext in sorted(chunker.supported_extensions):
        print(f"   • {ext}")
    
    print(f"\n⚙️  Configuration Options:")
    print(f"   • Chunk size: {chunker.chunk_size} tokens (configurable)")
    print(f"   • Chunk overlap: {chunker.chunk_overlap} tokens (configurable)")
    print(f"   • Token estimation: {chunker.chars_per_token} chars per token")
    
    print(f"\n🔍 Parsing Features:")
    print(f"   • PDF: Advanced text extraction with pdfplumber")
    print(f"   • DOCX: Text and table extraction with python-docx")
    print(f"   • CSV/Excel: Data conversion to readable text with pandas")
    print(f"   • Intelligent sentence-based chunking")
    print(f"   • Context preservation with overlap")
    
    print(f"\n📋 Metadata Tracking:")
    metadata_fields = [
        "source_file", "file_name", "file_hash", "page_number",
        "total_pages", "chunk_index", "chunk_type", "token_count",
        "char_count", "created_at", "chunk_size", "chunk_overlap"
    ]
    for field in metadata_fields:
        print(f"   • {field}")
    
    print(f"\n🚀 Integration Features:")
    print(f"   • FastAPI-ready functions")
    print(f"   • Batch processing support")
    print(f"   • JSON serialization")
    print(f"   • Comprehensive error handling")
    print(f"   • Structured logging")


if __name__ == "__main__":
    demonstrate_chunker()
    show_chunker_capabilities()
