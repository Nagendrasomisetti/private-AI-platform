"""
File Parser & Chunker for PrivAI
Handles parsing of various file types and intelligent text chunking for embeddings
"""
import os
import re
import time
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional, Union, Tuple
import hashlib

# Optional imports for file processing
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .logging import get_logger

logger = get_logger("chunker")


class FileChunker:
    """
    A comprehensive file parser and chunker that handles multiple file types
    and creates intelligent text chunks suitable for embeddings.
    """
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        """
        Initialize the FileChunker with configurable chunk parameters.
        
        Args:
            chunk_size: Target number of tokens per chunk (default: 500)
            chunk_overlap: Number of tokens to overlap between chunks (default: 50)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_extensions = {'.pdf', '.docx', '.doc', '.csv', '.xlsx', '.xls'}
        
        # Token counting approximation (roughly 4 characters per token)
        self.chars_per_token = 4
        
        logger.info("FileChunker initialized", 
                   chunk_size=chunk_size, 
                   chunk_overlap=chunk_overlap)
    
    def parse_file(self, file_path: Union[str, Path]) -> List[Dict[str, Any]]:
        """
        Parse a file and return chunks with metadata.
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            List of dictionaries containing chunk data and metadata
            
        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file doesn't exist
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        logger.info("Parsing file", file_path=str(file_path), extension=file_extension)
        
        # Generate file hash for tracking
        file_hash = self._generate_file_hash(file_path)
        
        # Parse based on file type
        if file_extension == '.pdf':
            return self._parse_pdf(file_path, file_hash)
        elif file_extension in ['.docx', '.doc']:
            return self._parse_docx(file_path, file_hash)
        elif file_extension in ['.csv', '.xlsx', '.xls']:
            return self._parse_tabular(file_path, file_hash)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _parse_pdf(self, file_path: Path, file_hash: str) -> List[Dict[str, Any]]:
        """
        Parse PDF file using pdfplumber for better text extraction.
        
        Args:
            file_path: Path to PDF file
            file_hash: File hash for tracking
            
        Returns:
            List of chunk dictionaries
        """
        if not PDFPLUMBER_AVAILABLE:
            raise ImportError("pdfplumber is required for PDF processing. Install with: pip install pdfplumber")
        
        chunks = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                logger.info("PDF parsing started", pages=total_pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text from page
                    page_text = page.extract_text()
                    
                    if not page_text or not page_text.strip():
                        logger.warning("Empty page found", page=page_num)
                        continue
                    
                    # Extract tables if present
                    tables = page.extract_tables()
                    table_text = ""
                    if tables:
                        table_text = self._format_tables(tables)
                    
                    # Combine page text and table text
                    full_text = page_text
                    if table_text:
                        full_text += "\n\nTables:\n" + table_text
                    
                    # Create chunks from page text
                    page_chunks = self._create_chunks(
                        text=full_text,
                        source_file=str(file_path),
                        file_hash=file_hash,
                        page_number=page_num,
                        total_pages=total_pages,
                        chunk_type="pdf_page"
                    )
                    
                    chunks.extend(page_chunks)
                
                logger.info("PDF parsing completed", 
                           file=str(file_path),
                           pages=total_pages,
                           chunks=len(chunks))
                
        except Exception as e:
            logger.error("PDF parsing failed", file=str(file_path), error=str(e))
            raise
        
        return chunks
    
    def _parse_docx(self, file_path: Path, file_hash: str) -> List[Dict[str, Any]]:
        """
        Parse DOCX file using python-docx.
        
        Args:
            file_path: Path to DOCX file
            file_hash: File hash for tracking
            
        Returns:
            List of chunk dictionaries
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
        
        chunks = []
        
        try:
            doc = Document(file_path)
            logger.info("DOCX parsing started", file=str(file_path))
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                table_text = self._format_docx_table(table)
                if table_text.strip():
                    table_texts.append(table_text)
            
            # Combine all text
            full_text = "\n\n".join(paragraphs)
            if table_texts:
                full_text += "\n\nTables:\n" + "\n\n".join(table_texts)
            
            # Create chunks
            chunks = self._create_chunks(
                text=full_text,
                source_file=str(file_path),
                file_hash=file_hash,
                page_number=1,  # DOCX doesn't have page numbers
                total_pages=1,
                chunk_type="docx_document"
            )
            
            logger.info("DOCX parsing completed", 
                       file=str(file_path),
                       chunks=len(chunks))
            
        except Exception as e:
            logger.error("DOCX parsing failed", file=str(file_path), error=str(e))
            raise
        
        return chunks
    
    def _parse_tabular(self, file_path: Path, file_hash: str) -> List[Dict[str, Any]]:
        """
        Parse CSV/Excel files using pandas.
        
        Args:
            file_path: Path to tabular file
            file_hash: File hash for tracking
            
        Returns:
            List of chunk dictionaries
        """
        if not PANDAS_AVAILABLE:
            raise ImportError("pandas is required for CSV/Excel processing. Install with: pip install pandas openpyxl")
        
        chunks = []
        
        try:
            file_extension = file_path.suffix.lower()
            logger.info("Tabular file parsing started", 
                       file=str(file_path), 
                       extension=file_extension)
            
            # Read file based on extension
            if file_extension == '.csv':
                df = pd.read_csv(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
            else:
                raise ValueError(f"Unsupported tabular file type: {file_extension}")
            
            # Convert DataFrame to text
            text_content = self._dataframe_to_text(df, str(file_path))
            
            # Create chunks
            chunks = self._create_chunks(
                text=text_content,
                source_file=str(file_path),
                file_hash=file_hash,
                page_number=1,
                total_pages=1,
                chunk_type="tabular_data"
            )
            
            logger.info("Tabular file parsing completed", 
                       file=str(file_path),
                       rows=len(df),
                       columns=len(df.columns),
                       chunks=len(chunks))
            
        except Exception as e:
            logger.error("Tabular file parsing failed", file=str(file_path), error=str(e))
            raise
        
        return chunks
    
    def _create_chunks(self, text: str, source_file: str, file_hash: str, 
                      page_number: int, total_pages: int, chunk_type: str) -> List[Dict[str, Any]]:
        """
        Create text chunks with metadata from input text.
        
        Args:
            text: Input text to chunk
            source_file: Source file path
            file_hash: File hash for tracking
            page_number: Page number (1-indexed)
            total_pages: Total pages in document
            chunk_type: Type of chunk (pdf_page, docx_document, etc.)
            
        Returns:
            List of chunk dictionaries
        """
        if not text or not text.strip():
            return []
        
        # Clean and normalize text
        cleaned_text = self._clean_text(text)
        
        # Split into sentences for better chunking
        sentences = self._split_into_sentences(cleaned_text)
        
        # Create chunks
        chunks = []
        current_chunk = ""
        current_tokens = 0
        chunk_index = 0
        
        for sentence in sentences:
            sentence_tokens = self._estimate_tokens(sentence)
            
            # If adding this sentence would exceed chunk size, finalize current chunk
            if current_tokens + sentence_tokens > self.chunk_size and current_chunk:
                chunk_data = self._create_chunk_metadata(
                    text=current_chunk.strip(),
                    source_file=source_file,
                    file_hash=file_hash,
                    page_number=page_number,
                    total_pages=total_pages,
                    chunk_index=chunk_index,
                    chunk_type=chunk_type,
                    total_tokens=current_tokens
                )
                chunks.append(chunk_data)
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + " " + sentence
                current_tokens = self._estimate_tokens(current_chunk)
                chunk_index += 1
            else:
                # Add sentence to current chunk
                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence
                current_tokens += sentence_tokens
        
        # Add final chunk if it has content
        if current_chunk.strip():
            chunk_data = self._create_chunk_metadata(
                text=current_chunk.strip(),
                source_file=source_file,
                file_hash=file_hash,
                page_number=page_number,
                total_pages=total_pages,
                chunk_index=chunk_index,
                chunk_type=chunk_type,
                total_tokens=current_tokens
            )
            chunks.append(chunk_data)
        
        logger.info("Chunks created", 
                   source_file=source_file,
                   total_chunks=len(chunks),
                   avg_tokens=sum(c['metadata']['token_count'] for c in chunks) // len(chunks) if chunks else 0)
        
        return chunks
    
    def _create_chunk_metadata(self, text: str, source_file: str, file_hash: str,
                              page_number: int, total_pages: int, chunk_index: int,
                              chunk_type: str, total_tokens: int) -> Dict[str, Any]:
        """
        Create metadata dictionary for a chunk.
        
        Args:
            text: Chunk text content
            source_file: Source file path
            file_hash: File hash for tracking
            page_number: Page number
            total_pages: Total pages
            chunk_index: Index of chunk within document
            chunk_type: Type of chunk
            total_tokens: Estimated token count
            
        Returns:
            Dictionary containing chunk data and metadata
        """
        timestamp = datetime.now().isoformat()
        
        return {
            "text": text,
            "metadata": {
                "source_file": source_file,
                "file_name": Path(source_file).name,
                "file_hash": file_hash,
                "page_number": page_number,
                "total_pages": total_pages,
                "chunk_index": chunk_index,
                "chunk_type": chunk_type,
                "token_count": total_tokens,
                "char_count": len(text),
                "created_at": timestamp,
                "chunk_size": self.chunk_size,
                "chunk_overlap": self.chunk_overlap
            }
        }
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize text for better chunking.
        
        Args:
            text: Raw text to clean
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}\"\'\/]', '', text)
        
        # Normalize quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        return text.strip()
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """
        Split text into sentences for better chunking.
        
        Args:
            text: Text to split
            
        Returns:
            List of sentences
        """
        # Simple sentence splitting (can be enhanced with NLTK)
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        # Filter out empty sentences
        sentences = [s.strip() for s in sentences if s.strip()]
        
        return sentences
    
    def _estimate_tokens(self, text: str) -> int:
        """
        Estimate token count for text (rough approximation).
        
        Args:
            text: Text to estimate
            
        Returns:
            Estimated token count
        """
        return len(text) // self.chars_per_token
    
    def _get_overlap_text(self, text: str) -> str:
        """
        Get overlap text from the end of current chunk.
        
        Args:
            text: Current chunk text
            
        Returns:
            Overlap text
        """
        words = text.split()
        overlap_words = words[-self.chunk_overlap:] if len(words) > self.chunk_overlap else words
        return " ".join(overlap_words)
    
    def _generate_file_hash(self, file_path: Path) -> str:
        """
        Generate MD5 hash for file tracking.
        
        Args:
            file_path: Path to file
            
        Returns:
            MD5 hash string
        """
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def _format_tables(self, tables: List[List[List[str]]]) -> str:
        """
        Format extracted tables into readable text.
        
        Args:
            tables: List of tables from PDF
            
        Returns:
            Formatted table text
        """
        formatted_tables = []
        
        for table in tables:
            if not table:
                continue
            
            table_text = "Table:\n"
            for row in table:
                if row:
                    table_text += " | ".join(str(cell) if cell else "" for cell in row) + "\n"
            formatted_tables.append(table_text)
        
        return "\n\n".join(formatted_tables)
    
    def _format_docx_table(self, table: Table) -> str:
        """
        Format DOCX table into readable text.
        
        Args:
            table: DOCX table object
            
        Returns:
            Formatted table text
        """
        table_text = "Table:\n"
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_text.append(cell_text if cell_text else "")
            table_text += " | ".join(row_text) + "\n"
        
        return table_text
    
    def _dataframe_to_text(self, df, source_file: str) -> str:
        """
        Convert pandas DataFrame to readable text format.
        
        Args:
            df: DataFrame to convert
            source_file: Source file path for context
            
        Returns:
            Formatted text representation
        """
        text_parts = []
        
        # Add file context
        text_parts.append(f"Data from: {Path(source_file).name}")
        text_parts.append(f"Shape: {df.shape[0]} rows, {df.shape[1]} columns")
        text_parts.append("")
        
        # Add column information
        text_parts.append("Columns:")
        for i, col in enumerate(df.columns):
            dtype = str(df[col].dtype)
            text_parts.append(f"  {i+1}. {col} ({dtype})")
        text_parts.append("")
        
        # Add sample data
        text_parts.append("Sample data:")
        sample_df = df.head(10)  # First 10 rows
        
        for idx, row in sample_df.iterrows():
            row_text = f"Row {idx + 1}: "
            for col in df.columns:
                value = str(row[col]) if pd.notna(row[col]) else "N/A"
                row_text += f"{col}: {value}, "
            text_parts.append(row_text.rstrip(", "))
        
        # Add summary statistics for numeric columns
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            text_parts.append("")
            text_parts.append("Summary statistics:")
            for col in numeric_cols:
                stats = df[col].describe()
                text_parts.append(f"  {col}: mean={stats['mean']:.2f}, std={stats['std']:.2f}, min={stats['min']:.2f}, max={stats['max']:.2f}")
        
        return "\n".join(text_parts)
    
    def get_chunking_stats(self, chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Get statistics about chunking results.
        
        Args:
            chunks: List of chunk dictionaries
            
        Returns:
            Dictionary with chunking statistics
        """
        if not chunks:
            return {"total_chunks": 0}
        
        token_counts = [chunk['metadata']['token_count'] for chunk in chunks]
        
        return {
            "total_chunks": len(chunks),
            "avg_tokens_per_chunk": sum(token_counts) / len(token_counts),
            "min_tokens": min(token_counts),
            "max_tokens": max(token_counts),
            "total_tokens": sum(token_counts),
            "chunk_size_target": self.chunk_size,
            "chunk_overlap": self.chunk_overlap
        }


# Convenience functions for FastAPI integration
def parse_file_to_chunks(file_path: Union[str, Path], 
                        chunk_size: int = 500, 
                        chunk_overlap: int = 50) -> List[Dict[str, Any]]:
    """
    Convenience function to parse a file and return chunks.
    
    Args:
        file_path: Path to file to parse
        chunk_size: Target tokens per chunk
        chunk_overlap: Token overlap between chunks
        
    Returns:
        List of chunk dictionaries
    """
    chunker = FileChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return chunker.parse_file(file_path)


def parse_multiple_files(file_paths: List[Union[str, Path]], 
                        chunk_size: int = 500, 
                        chunk_overlap: int = 50) -> List[Dict[str, Any]]:
    """
    Parse multiple files and return all chunks.
    
    Args:
        file_paths: List of file paths to parse
        chunk_size: Target tokens per chunk
        chunk_overlap: Token overlap between chunks
        
    Returns:
        List of all chunk dictionaries from all files
    """
    chunker = FileChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    all_chunks = []
    
    for file_path in file_paths:
        try:
            chunks = chunker.parse_file(file_path)
            all_chunks.extend(chunks)
            logger.info("File parsed successfully", 
                       file=str(file_path), 
                       chunks=len(chunks))
        except Exception as e:
            logger.error("Failed to parse file", 
                        file=str(file_path), 
                        error=str(e))
    
    return all_chunks


# Global chunker instance for reuse
default_chunker = FileChunker()
